#!/usr/bin/env python3
"""Convert MES edge AI paper from MD to DOCX"""
import os
import re
import hashlib
import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

# OMML namespace
NS_M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
def _m(tag):
    return f'{{{NS_M}}}{tag}'

# LaTeX command -> Unicode mapping
LATEX_UNICODE = {
    r'\longrightarrow': '\u2192',
    r'\rightarrow': '\u2192',
    r'\Rightarrow': '\u21d2',
    r'\leftarrow': '\u2190',
    r'\Leftarrow': '\u21d0',
    r'\xrightarrow': '',  # handled specially
    r'\times': '\u00d7',
    r'\cdot': '\u00b7',
    r'\pm': '\u00b1',
    r'\infty': '\u221e',
    r'\partial': '\u2202',
    r'\nabla': '\u2207',
    r'\approx': '\u2248',
    r'\neq': '\u2260',
    r'\leq': '\u2264',
    r'\geq': '\u2265',
    r'\subset': '\u2282',
    r'\supset': '\u2283',
    r'\subseteq': '\u2286',
    r'\supseteq': '\u2287',
    r'\cup': '\u222a',
    r'\cap': '\u2229',
    r'\in': '\u2208',
    r'\notin': '\u2209',
    r'\forall': '\u2200',
    r'\exists': '\u2203',
    r'\emptyset': '\u2205',
    r'\alpha': '\u03b1',
    r'\beta': '\u03b2',
    r'\gamma': '\u03b3',
    r'\delta': '\u03b4',
    r'\epsilon': '\u03b5',
    r'\zeta': '\u03b6',
    r'\eta': '\u03b7',
    r'\theta': '\u03b8',
    r'\lambda': '\u03bb',
    r'\mu': '\u03bc',
    r'\sigma': '\u03c3',
    r'\tau': '\u03c4',
    r'\phi': '\u03c6',
    r'\omega': '\u03c9',
    r'\Delta': '\u0394',
    r'\Sigma': '\u03a3',
    r'\Theta': '\u0398',
    r'\Omega': '\u03a9',
    r'\%': '%',
    r'\bar': '',  # handled in context
}

# Greek letters that need italic OMML runs (not in LATEX_UNICODE)
GREEK_CMDS = {
    r'\alpha', r'\beta', r'\gamma', r'\delta', r'\epsilon',
    r'\zeta', r'\eta', r'\theta', r'\lambda', r'\mu',
    r'\sigma', r'\tau', r'\phi', r'\omega', r'\Delta',
    r'\Sigma', r'\Theta', r'\Omega',
}


# ---------------------------------------------------------------------------
# LaTeX math -> OMML (Word native editable formulas)
# ---------------------------------------------------------------------------

def _latex_tokenize(text):
    """Tokenize LaTeX math text into a list of tokens."""
    tokens = []
    i = 0
    while i < len(text):
        if text[i] in ' \t':
            i += 1
            continue
        c = text[i]
        if c == '\\':
            j = i + 1
            while j < len(text) and text[j].isalpha():
                j += 1
            tokens.append(('cmd', text[i:j]))
            i = j
        elif c == '{':
            depth = 1
            j = i + 1
            while j < len(text) and depth > 0:
                if text[j] == '{':
                    depth += 1
                elif text[j] == '}':
                    depth -= 1
                j += 1
            tokens.append(('group', text[i+1:j-1]))
            i = j
        elif c == '}':
            tokens.append(('close', '}'))
            i += 1
        elif c == '_':
            tokens.append(('sub', '_'))
            i += 1
        elif c == '^':
            tokens.append(('sup', '^'))
            i += 1
        else:
            j = i
            while j < len(text) and text[j] not in '\\_{} \t':
                j += 1
            tokens.append(('text', text[i:j]))
            i = j
    return tokens


def _read_group_or_text(tokens, pos):
    """Read next token which can be a group or plain text. Returns (content_string, new_pos)."""
    if pos >= len(tokens):
        return '', pos
    tok_type, tok_val = tokens[pos]
    if tok_type == 'group':
        return tok_val, pos + 1
    if tok_type == 'text':
        return tok_val, pos + 1
    return '', pos


def _build_omml(math_elem, tokens, pos=0):
    """
    Parse LaTeX tokens and build OMML sub-elements under math_elem.
    Returns the new position after parsing.
    """
    while pos < len(tokens):
        tok_type, tok_val = tokens[pos]

        if tok_type == 'close':
            break

        if tok_type == 'sub':
            # Subscript: read next token as the subscript
            pos += 1
            sub_content, pos = _read_group_or_text(tokens, pos)
            # Insert as subscript
            _add_omml_subscript(math_elem, None, sub_content)
            continue

        if tok_type == 'sup':
            pos += 1
            sup_content, pos = _read_group_or_text(tokens, pos)
            _add_omml_superscript(math_elem, None, sup_content)
            continue

        if tok_type == 'group':
            # Parse and insert group content
            group_math = etree.SubElement(math_elem, _m('d'))
            pos = _build_omml(group_math, _latex_tokenize(tok_val))
            continue

        if tok_type == 'text':
            _add_omml_text_run(math_elem, tok_val)
            pos += 1
            continue

        if tok_type == 'cmd':
            cmd = tok_val
            pos += 1

            if cmd == r'\text':
                content, pos = _read_group_or_text(tokens, pos)
                _add_omml_text_run(math_elem, content)

            elif cmd == r'\mathbf':
                content, pos = _read_group_or_text(tokens, pos)
                _add_omml_text_run(math_elem, content, bold=True)

            elif cmd == r'\frac':
                num, pos = _read_group_or_text(tokens, pos)
                den, pos = _read_group_or_text(tokens, pos)
                _add_omml_fraction(math_elem, num, den)

            elif cmd == r'\sqrt':
                content, pos = _read_group_or_text(tokens, pos)
                _add_omml_radical(math_elem, content)

            elif cmd in (r'\sum', r'\prod', r'\int'):
                upper = ''
                lower = ''
                for _ in range(2):
                    if pos < len(tokens):
                        nxt_type, _ = tokens[pos]
                        if nxt_type == 'sup':
                            pos += 1
                            upper, pos = _read_group_or_text(tokens, pos)
                        elif nxt_type == 'sub':
                            pos += 1
                            lower, pos = _read_group_or_text(tokens, pos)
                        else:
                            break
                _add_omml_nary(math_elem, cmd, lower, upper)

            elif cmd == r'\xrightarrow':
                content, pos = _read_group_or_text(tokens, pos)
                _add_omml_text_run(math_elem, '\u2014' + content + '\u2192')

            elif cmd == r'\longrightarrow':
                _add_omml_text_run(math_elem, '\u2192')

            elif cmd == r'\rightarrow':
                _add_omml_text_run(math_elem, '\u2192')

            elif cmd == r'\Rightarrow':
                _add_omml_text_run(math_elem, '\u21d2')

            elif cmd == r'\leftarrow':
                _add_omml_text_run(math_elem, '\u2190')

            elif cmd == r'\Leftarrow':
                _add_omml_text_run(math_elem, '\u21d0')

            elif cmd in GREEK_CMDS:
                unicode_char = LATEX_UNICODE.get(cmd, '')
                if unicode_char:
                    _add_omml_text_run(math_elem, unicode_char)

            elif cmd == r'\%':
                _add_omml_text_run(math_elem, '%')

            elif cmd in LATEX_UNICODE:
                unicode_char = LATEX_UNICODE[cmd]
                if unicode_char:
                    _add_omml_text_run(math_elem, unicode_char)

            elif cmd == r'\bar':
                content, pos = _read_group_or_text(tokens, pos)
                _add_omml_bar(math_elem, content)

            elif cmd == r'\mathrm':
                content, pos = _read_group_or_text(tokens, pos)
                _add_omml_text_run(math_elem, content)

            elif cmd == r'\textit':
                content, pos = _read_group_or_text(tokens, pos)
                _add_omml_text_run(math_elem, content)

            else:
                # Unknown command - try to read its argument if it has one
                if pos < len(tokens) and tokens[pos][0] == 'group':
                    # Just skip the argument for unknown commands
                    _, pos = _read_group_or_text(tokens, pos)

    return pos


def _add_omml_text_run(elem, text, bold=False):
    """Add an m:r (run) with m:t (text) to the OMML element."""
    if not text:
        text = ' '
    r = etree.SubElement(elem, _m('r'))
    rPr = etree.SubElement(r, _m('rPr'))
    # Make variables (word chars) italic by default, numbers/digits not
    style = etree.SubElement(rPr, _m('sty'))
    style.text = 'p'  # plain style for text, no italics by default
    # Add script font
    lc = etree.SubElement(rPr, _m('lit'))
    lc.text = '0'
    t = etree.SubElement(r, _m('t'))
    t.text = text
    t.set(f'{{http://www.w3.org/XML/1998/namespace}}space', 'preserve')
    return r


def _add_omml_fraction(elem, num_latex, den_latex):
    """Add m:f (fraction) element."""
    f = etree.SubElement(elem, _m('f'))
    num = etree.SubElement(f, _m('num'))
    _build_omml(num, _latex_tokenize(num_latex))
    den = etree.SubElement(f, _m('den'))
    _build_omml(den, _latex_tokenize(den_latex))
    return f


def _add_omml_subscript(elem, base_latex, sub_latex):
    """Add m:sSub (subscript) element."""
    ssub = etree.SubElement(elem, _m('sSub'))
    if base_latex:
        e = etree.SubElement(ssub, _m('e'))
        _build_omml(e, _latex_tokenize(base_latex))
    sub_e = etree.SubElement(ssub, _m('sub'))
    _build_omml(sub_e, _latex_tokenize(sub_latex))
    return ssub


def _add_omml_superscript(elem, base_latex, sup_latex):
    """Add m:sSup (superscript) element."""
    ssup = etree.SubElement(elem, _m('sSup'))
    if base_latex:
        e = etree.SubElement(ssup, _m('e'))
        _build_omml(e, _latex_tokenize(base_latex))
    sup_e = etree.SubElement(ssup, _m('sup'))
    _build_omml(sup_e, _latex_tokenize(sup_latex))
    return ssup


def _add_omml_nary(elem, cmd, lower_latex, upper_latex):
    """Add m:nary (n-ary operator like sum/prod/int)."""
    op_map = {
        r'\sum': '\u2211',
        r'\prod': '\u220f',
        r'\int': '\u222b',
    }
    nary = etree.SubElement(elem, _m('nary'))
    naryPr = etree.SubElement(nary, _m('naryPr'))
    chr_e = etree.SubElement(naryPr, _m('chr'))
    chr_e.text = op_map.get(cmd, '\u2211')
    if lower_latex:
        sub_e = etree.SubElement(nary, _m('sub'))
        _build_omml(sub_e, _latex_tokenize(lower_latex))
    if upper_latex:
        sup_e = etree.SubElement(nary, _m('sup'))
        _build_omml(sup_e, _latex_tokenize(upper_latex))
    return nary


def _add_omml_radical(elem, content):
    """Add m:rad (radical/square root)."""
    rad = etree.SubElement(elem, _m('rad'))
    radPr = etree.SubElement(rad, _m('radPr'))
    deg_hide = etree.SubElement(radPr, _m('degHide'))
    deg_hide.set(_m('val'), '1')
    e = etree.SubElement(rad, _m('e'))
    _build_omml(e, _latex_tokenize(content))
    return rad


def _add_omml_bar(elem, content):
    """Add m:bar (overline/accent bar)."""
    bar = etree.SubElement(elem, _m('bar'))
    barPr = etree.SubElement(bar, _m('barPr'))
    pos = etree.SubElement(barPr, _m('pos'))
    pos.set(_m('val'), 'top')
    e = etree.SubElement(bar, _m('e'))
    _build_omml(e, _latex_tokenize(content))
    return bar


def latex_to_omml_para(latex_text):
    """
    Convert a LaTeX math expression into an OMML m:oMathPara XML element.
    Returns an lxml Element ready to be appended to a paragraph.
    """
    para = etree.Element(_m('oMathPara'))
    math = etree.SubElement(para, _m('oMath'))
    tokens = _latex_tokenize(latex_text)
    _build_omml(math, tokens)
    return para


def _add_center_text(doc, text):
    """Helper: add a centered plain-text paragraph (fallback)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, '宋体', '宋体', Pt(12))
    return p


def _handle_cases_formula(doc, formula):
    """Handle \begin{cases}...\end{cases} with per-line formulas."""
    match = re.search(r'\\begin\{cases\}(.*?)\\end\{cases\}', formula, re.DOTALL)
    if match:
        cases_content = match.group(1)
        for case_line in cases_content.split(r'\\'):
            case_line = case_line.strip().strip(',')
            if case_line:
                add_omml_formula(doc, case_line)

    remaining = formula.split(r'\end{cases}')[-1].strip()
    if remaining:
        add_omml_formula(doc, remaining)

def _handle_aligned_formula(doc, formula):
    """Handle \begin{aligned} with per-line formulas."""
    match = re.search(r'\\begin\{aligned\}(.*?)\\end\{aligned\}', formula, re.DOTALL)
    if match:
        content = match.group(1)
        for line in content.split(r'\\'):
            line = line.strip().replace('&', '').strip()
            if line:
                add_omml_formula(doc, line)

def render_formula_to_png(formula_text, output_path, fontsize=14, dpi=200):
    """Render LaTeX formula text as PNG image using matplotlib."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.font_manager import FontProperties

    # Find a CJK-supporting font for Chinese characters in formulas
    _CJK_FONT_CANDIDATES = [
        '/mnt/c/Windows/Fonts/msyh.ttc',       # Microsoft YaHei
        '/mnt/c/Windows/Fonts/simhei.ttf',     # SimHei
        '/mnt/c/Windows/Fonts/simsun.ttc',     # SimSun
        '/mnt/c/Windows/Fonts/yahei.ttf',      # YaHei (alt)
    ]
    cjk_font_path = None
    for fp in _CJK_FONT_CANDIDATES:
        if os.path.exists(fp):
            cjk_font_path = fp
            break

    text = formula_text

    # Apply Unicode replacements for common LaTeX commands
    sorted_cmds = sorted(LATEX_UNICODE.items(), key=lambda x: -len(x[0]))
    for cmd, unicode_char in sorted_cmds:
        text = text.replace(cmd, unicode_char)

    # Remove remaining LaTeX commands (backslash-word)
    text = re.sub(r'\\[a-zA-Z]+', '', text)

    # Clean braces
    text = text.replace('{', '').replace('}', '')

    # Remove \begin{xxx}, \end{xxx} environment markers
    text = re.sub(r'\\begin\{[a-zA-Z]+\}', '', text)
    text = re.sub(r'\\end\{[a-zA-Z]+\}', '', text)

    # Replace \\ with newline for multi-line formulas
    text = text.replace(r'\\', '\n')

    # Remove alignment markers
    text = text.replace('&', '')

    # Clean up whitespace
    text = re.sub(r'\n\s+', '\n', text)
    text = text.strip()

    if not text:
        text = ' '

    # Calculate figure size based on content
    num_lines = text.count('\n') + 1
    max_line_len = max(len(l) for l in text.split('\n')) if text else 1

    fig_width = min(max_line_len * 0.12, 5.5)
    fig_height = 0.6 + num_lines * 0.45

    fig, ax = plt.subplots(figsize=(fig_width, fig_height))

    if cjk_font_path:
        fp = FontProperties(fname=cjk_font_path, size=fontsize)
        ax.text(0.5, 0.5, text, fontproperties=fp, ha='center', va='center',
                transform=ax.transAxes)
    else:
        ax.text(0.5, 0.5, text, fontsize=fontsize, ha='center', va='center',
                fontfamily='serif', transform=ax.transAxes)

    ax.axis('off')
    plt.savefig(output_path, dpi=dpi, bbox_inches='tight',
                facecolor='white', pad_inches=0.15)
    plt.close()


def add_omml_formula(doc, formula_text):
    """
    Render formula as PNG image and insert into document.
    """
    formula = formula_text.strip()

    if not formula:
        return

    # Multi-line environments — render whole formula as one image
    formula_hash = hashlib.md5(formula.encode()).hexdigest()[:12]
    img_path = f'/tmp/formula_{formula_hash}.png'

    try:
        render_formula_to_png(formula, img_path)
        doc.add_picture(img_path, width=Inches(5.0))
    except Exception as e:
        print(f"\u26a0\ufe0f Formula render failed, falling back to text: {e}")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(formula)
        set_run_font(run, '宋体', '宋体', Pt(12))
    finally:
        if os.path.exists(img_path):
            os.remove(img_path)


def create_document():
    """Create document with proper page setup"""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.95)
    section.bottom_margin = Inches(0.95)
    section.left_margin = Inches(1.02)
    section.right_margin = Inches(1.02)

    return doc


def set_run_font(run, font_name_ascii='宋体', font_name_east='宋体', size=Pt(12), bold=False):
    """Set font for a run"""
    run.font.name = font_name_ascii
    run.font.size = size
    run.font.bold = bold
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_east)
    rFonts.set(qn('w:ascii'), font_name_ascii)
    rFonts.set(qn('w:hAnsi'), font_name_ascii)


def set_paragraph_spacing(paragraph, line_spacing=Pt(20), before=0, after=0, first_line_indent=None):
    """Set paragraph spacing"""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = line_spacing
    pf.space_before = before
    pf.space_after = after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def add_heading_text(doc, text, level=1):
    """Add heading with proper formatting"""
    p = doc.add_paragraph()

    if level == 0:  # Document title (大标题)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, Pt(20), before=0, after=Pt(30))
        run = p.add_run(text)
        set_run_font(run, '黑体', '黑体', Pt(15), bold=True)
    elif level == 1:  # Chapter (## 第X章)
        set_paragraph_spacing(p, Pt(20), before=0, after=Pt(18))
        run = p.add_run(text)
        set_run_font(run, '黑体', '黑体', Pt(14), bold=True)
    elif level == 2:  # Section (### X.Y)
        set_paragraph_spacing(p, Pt(20), before=0, after=Pt(12))
        run = p.add_run(text)
        set_run_font(run, '黑体', '黑体', Pt(14), bold=True)
    elif level == 3:  # Subsection (#### X.Y.Z)
        set_paragraph_spacing(p, Pt(20), before=0, after=Pt(6))
        run = p.add_run(text)
        set_run_font(run, '黑体', '黑体', Pt(12), bold=True)

    return p


def add_body_text(doc, text, indent=True):
    """Add body paragraph with 宋体 小四"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, Pt(20), before=0, after=0, first_line_indent=Pt(24) if indent else None)
    run = p.add_run(text)
    set_run_font(run, '宋体', '宋体', Pt(12))
    return p


def parse_md_table(lines):
    """Parse pipe table lines into rows/cols"""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith('|-'):  # skip separator
            continue
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.strip('|').split('|')]
            rows.append(cells)
    return rows


def add_table(doc, rows):
    """Add a docx table from parsed rows"""
    if not rows:
        return
    nrows = len(rows)
    ncols = max(len(r) for r in rows)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style
    table.style = 'Table Grid'

    for ri, row_data in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.text = row_data[ci] if ci < len(row_data) else ''
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, '宋体', '宋体', Pt(10.5), bold=(ri == 0))

    return table


def _insert_chart_image(doc, img_path, width_inches=4.5):
    """Insert a chart image centered, with figure label."""
    if not os.path.exists(img_path):
        print(f'  ⚠️  Image not found: {img_path}')
        return
    doc.add_picture(img_path, width=Inches(width_inches))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = create_document()
    lines = md_text.split('\n')

    i = 0
    in_code_block = False
    in_table = False
    table_lines = []
    _inserted = set()  # track which chart images have been inserted

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, Pt(12), before=0, after=0)
            run = p.add_run(line)
            set_run_font(run, 'Consolas', '仿宋', Pt(9))
            i += 1
            continue

        # Pipe table - collect consecutive table lines
        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_lines.append(line)
            in_table = True
            i += 1
            # Check next line
            continue

        # If we were in a table, flush it
        if in_table:
            rows = parse_md_table(table_lines)
            if rows:
                doc.add_paragraph()  # blank line before
                add_table(doc, rows)
                doc.add_paragraph()  # blank line after
            table_lines = []
            in_table = False
            # Don't skip current line - process it below

        stripped = line.strip()

        # Skip empty lines (handled by table flush above)
        if not stripped:
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            add_heading_text(doc, text, level)

            # --- Chart image insertion triggers based on heading text ---
            heading_clean = re.sub(r'[\s\\/]', '', text)

            # 图1: 混淆矩阵 — insert after §3.3.4 heading
            if '混淆矩阵' in text and 'fig_cm' not in _inserted:
                _insert_chart_image(doc, '/mnt/d/tmp/fig_confusion_matrix.png')
                _inserted.add('fig_cm')

            # 图4: ST-GCN vs 规则分类器 — insert after 公开数据集验证 or 结果讨论
            if '公开数据集验证' in text and 'fig_model' not in _inserted:
                _insert_chart_image(doc, '/mnt/d/tmp/fig_model_comparison.png', 5.0)
                _inserted.add('fig_model')

            # 图3: 山积图 — insert after §5.4.2 heading
            if '山积图' in text and 'fig_lb' not in _inserted:
                _insert_chart_image(doc, '/mnt/d/tmp/fig_line_balance.png')
                _inserted.add('fig_lb')

            # 图2: 效率对比 — insert after §5.3.2 heading (效率分析)
            if '效率分析' in text and 'fig_eff' not in _inserted:
                _insert_chart_image(doc, '/mnt/d/tmp/fig_efficiency.png')
                _inserted.add('fig_eff')

            # 图5: KPI — insert after §5.4.1 heading (核心指标计算)
            if '核心指标计算' in text and 'fig_kpi' not in _inserted:
                _insert_chart_image(doc, '/mnt/d/tmp/fig_kpi.png')
                _inserted.add('fig_kpi')

            i += 1
            continue

        # Image placeholder
        img_match = re.match(r'!\[(.*?)\]\(.*?\)', stripped)
        if img_match:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'[图: {img_match.group(1)} - 待插入截图]')
            set_run_font(run, '宋体', '宋体', Pt(10.5))
            i += 1
            continue

        # Formula: single-line $$...$$
        if stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 2:
            formula = stripped.strip('$')
            add_omml_formula(doc, formula)
            i += 1
            continue

        # Formula: multi-line $$...$$ (opening $$)
        if stripped == '$$':
            # Collect lines until closing $$
            j = i + 1
            formula_lines = []
            while j < len(lines):
                if lines[j].strip() == '$$':
                    break
                formula_lines.append(lines[j])
                j += 1
            if j < len(lines):
                # Found closing $$
                formula_text = '\n'.join(formula_lines)
                if formula_text.strip():
                    doc.add_paragraph()  # blank line before formula
                    add_omml_formula(doc, formula_text)
                    doc.add_paragraph()  # blank line after
                i = j + 1  # skip past closing $$
                continue

        # Remove image from text references: ![text](url) → text
        text = re.sub(r'!\[(.*?)\]\(.*?\)', r'\1', stripped)
        # Remove regular links: [text](url) → text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        # Remove bold markers
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)

        if text:
            add_body_text(doc, text)

        i += 1

    doc.save(docx_path)
    print(f"✅ Saved to {docx_path}")
    return doc


if __name__ == '__main__':
    md_path = '/mnt/d/tmp/论文_边缘AI_优化完整版_v2.md'
    docx_path = '/mnt/d/tmp/论文_边缘AI_DOCX.docx'
    convert_md_to_docx(md_path, docx_path)
